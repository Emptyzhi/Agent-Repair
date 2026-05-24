"""Constrained tool-planning layer for Full-ours repair runs.

The planner is intentionally optional. It lets the model propose a small repair
plan, then executes only allowlisted, deterministic local inspection tools over
the current task resources. The resulting evidence cache can be injected into
the Full-ours generation prompt. Unsupported tool requests are recorded as
blocked evidence instead of being simulated.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any, Callable


TEXT_SUFFIXES = {".csv", ".json", ".md", ".txt", ".html", ".htm", ".xml", ".yaml", ".yml"}
MAX_TEXT_CHARS = 6000
MAX_ROWS = 12


def dump_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _task_resource_entries(task: dict[str, Any]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for item in task.get("files", []) or []:
        if isinstance(item, dict):
            entries.append(item)
    for dialog in task.get("dialogs", []) or []:
        for item in dialog.get("resources", []) or []:
            if isinstance(item, dict):
                entries.append(item)
            elif isinstance(item, str):
                entries.append({"path": item})
    return entries


def resource_map(task: dict[str, Any], dataset_dir: Path) -> dict[str, Path]:
    resources: dict[str, Path] = {}
    root = dataset_dir.resolve()
    for entry in _task_resource_entries(task):
        raw = entry.get("path") or entry.get("file") or entry.get("filename")
        if not raw:
            continue
        candidate = (root / str(raw)).resolve()
        try:
            candidate.relative_to(root)
        except ValueError:
            continue
        if candidate.exists() and candidate.is_file():
            resources[str(raw)] = candidate
            resources[candidate.name] = candidate
    return resources


def _resolve_resource(name: str, resources: dict[str, Path]) -> Path | None:
    if name in resources:
        return resources[name]
    lowered = name.lower()
    for key, path in resources.items():
        if key.lower() == lowered or path.name.lower() == lowered:
            return path
    return None


def _shorten(text: str, max_chars: int = MAX_TEXT_CHARS) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n...[truncated]..."


def build_tool_plan_prompt(
    *,
    task_id: int,
    task: dict[str, Any],
    diagnostics: list[dict[str, Any]],
    route: dict[str, Any],
    expected_files: list[str],
    available_resources: list[str],
) -> str:
    origin_prompt = ""
    for dialog in task.get("dialogs", []) or []:
        if dialog.get("role") == "user":
            origin_prompt = str(dialog.get("content", ""))
            break
    return f"""
You are planning a constrained repair for GTA-2 task {task_id}.

You may ask for local deterministic tool evidence before candidate generation.
Only these tools are available:
- list_task_resources: list declared task resource files.
- read_text_resource: read a declared text-like resource. Requires resource.
- summarize_csv_resource: inspect headers and sample rows of a declared CSV resource. Requires resource.
- extract_pdf_text: extract limited text from a declared PDF resource. Requires resource.
- extract_docx_text: extract limited text from a declared DOCX resource. Requires resource.
- inspect_image_resource: inspect dimensions and mode of a declared image resource. Requires resource.

Do not request web browsing, arbitrary shell commands, network APIs, hidden files, or paths not listed as task resources.
If the repair can be done without tool evidence, return an empty tool_plan.

Original task:
{origin_prompt}

Low-checkpoint diagnostics:
{json.dumps(diagnostics, ensure_ascii=False, indent=2)}

Frozen route metadata:
{json.dumps(route, ensure_ascii=False, indent=2)}

Expected deliverable filenames:
{json.dumps(expected_files, ensure_ascii=False, indent=2)}

Available declared resources:
{json.dumps(available_resources, ensure_ascii=False, indent=2)}

Return JSON only:
{{
  "failure_hypothesis": "short explanation",
  "tool_plan": [
    {{
      "tool": "summarize_csv_resource",
      "resource": "resources/files/example.csv",
      "purpose": "why this evidence is needed",
      "expected_evidence": "what should be extracted"
    }}
  ],
  "candidate_requirements": ["requirements the final candidate must satisfy"]
}}
""".strip()


def call_tool_planner(
    *,
    llm_json: Callable[[str, str, int], dict[str, Any]],
    prompt: str,
    model: str,
    max_tokens: int = 2048,
) -> dict[str, Any]:
    plan = llm_json(prompt, model, max_tokens)
    if not isinstance(plan, dict):
        return {"failure_hypothesis": "planner returned non-object", "tool_plan": [], "candidate_requirements": []}
    raw_steps = plan.get("tool_plan")
    if not isinstance(raw_steps, list):
        plan["tool_plan"] = []
    plan.setdefault("candidate_requirements", [])
    plan.setdefault("failure_hypothesis", "")
    return plan


def _list_task_resources(resources: dict[str, Path]) -> dict[str, Any]:
    rows = []
    seen: set[Path] = set()
    for label, path in sorted(resources.items()):
        if path in seen:
            continue
        seen.add(path)
        rows.append({"resource": label, "name": path.name, "suffix": path.suffix.lower(), "size_bytes": path.stat().st_size})
    return {"resources": rows}


def _read_text_resource(path: Path) -> dict[str, Any]:
    if path.suffix.lower() not in TEXT_SUFFIXES:
        return {"error": f"resource is not text-like: {path.suffix}"}
    return {"text": _shorten(path.read_text(encoding="utf-8-sig", errors="replace"))}


def _summarize_csv_resource(path: Path) -> dict[str, Any]:
    if path.suffix.lower() != ".csv":
        return {"error": f"resource is not csv: {path.suffix}"}
    with path.open("r", encoding="utf-8-sig", newline="", errors="replace") as f:
        reader = csv.DictReader(f)
        rows = []
        for index, row in enumerate(reader):
            if index >= MAX_ROWS:
                break
            rows.append(row)
        return {"headers": reader.fieldnames or [], "sample_rows": rows, "sample_size": len(rows)}


def _extract_pdf_text(path: Path) -> dict[str, Any]:
    if path.suffix.lower() != ".pdf":
        return {"error": f"resource is not pdf: {path.suffix}"}
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - depends on optional package
        return {"error": f"pypdf unavailable: {type(exc).__name__}"}
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages[:3]):
        pages.append({"page": index + 1, "text": _shorten(page.extract_text() or "", 2000)})
    return {"num_pages": len(reader.pages), "pages": pages}


def _extract_docx_text(path: Path) -> dict[str, Any]:
    if path.suffix.lower() != ".docx":
        return {"error": f"resource is not docx: {path.suffix}"}
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - depends on optional package
        return {"error": f"python-docx unavailable: {type(exc).__name__}"}
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return {"paragraph_count": len(doc.paragraphs), "text": _shorten(text)}


def _inspect_image_resource(path: Path) -> dict[str, Any]:
    if path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
        return {"error": f"resource is not image-like: {path.suffix}"}
    try:
        from PIL import Image
    except Exception as exc:  # pragma: no cover - depends on optional package
        return {"error": f"PIL unavailable: {type(exc).__name__}"}
    with Image.open(path) as img:
        return {"width": img.width, "height": img.height, "mode": img.mode, "format": img.format}


def execute_tool_plan(
    *,
    plan: dict[str, Any],
    task: dict[str, Any],
    dataset_dir: Path,
    out_dir: Path,
) -> dict[str, Any]:
    resources = resource_map(task, dataset_dir)
    steps = plan.get("tool_plan") if isinstance(plan.get("tool_plan"), list) else []
    evidence_steps = []
    for index, step in enumerate(steps[:8], 1):
        if not isinstance(step, dict):
            continue
        tool = str(step.get("tool") or "").strip()
        resource_name = str(step.get("resource") or "").strip()
        record: dict[str, Any] = {
            "index": index,
            "tool": tool,
            "resource": resource_name,
            "purpose": step.get("purpose", ""),
            "expected_evidence": step.get("expected_evidence", ""),
        }
        try:
            if tool == "list_task_resources":
                record["status"] = "ok"
                record["result"] = _list_task_resources(resources)
            else:
                path = _resolve_resource(resource_name, resources)
                if path is None:
                    record["status"] = "blocked"
                    record["result"] = {"error": "resource is not declared for this task"}
                elif tool == "read_text_resource":
                    record["status"] = "ok"
                    record["result"] = _read_text_resource(path)
                elif tool == "summarize_csv_resource":
                    record["status"] = "ok"
                    record["result"] = _summarize_csv_resource(path)
                elif tool == "extract_pdf_text":
                    record["status"] = "ok"
                    record["result"] = _extract_pdf_text(path)
                elif tool == "extract_docx_text":
                    record["status"] = "ok"
                    record["result"] = _extract_docx_text(path)
                elif tool == "inspect_image_resource":
                    record["status"] = "ok"
                    record["result"] = _inspect_image_resource(path)
                else:
                    record["status"] = "blocked"
                    record["result"] = {"error": "tool is not in allowlist"}
        except Exception as exc:
            record["status"] = "error"
            record["result"] = {"error": f"{type(exc).__name__}: {exc}"}
        evidence_steps.append(record)

    evidence = {
        "planner_enabled": True,
        "failure_hypothesis": plan.get("failure_hypothesis", ""),
        "candidate_requirements": plan.get("candidate_requirements", []),
        "available_resources": sorted(set(resources)),
        "steps": evidence_steps,
    }
    dump_json(out_dir / "tool_repair_plan.json", plan)
    dump_json(out_dir / "tool_evidence_cache.json", evidence)
    return evidence


def format_evidence_for_prompt(evidence: dict[str, Any] | None) -> str:
    if not evidence:
        return "No tool evidence cache was requested or available."
    compact = {
        "failure_hypothesis": evidence.get("failure_hypothesis", ""),
        "candidate_requirements": evidence.get("candidate_requirements", []),
        "steps": evidence.get("steps", []),
    }
    return json.dumps(compact, ensure_ascii=False, indent=2)
